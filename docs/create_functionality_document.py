from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_word_document():
    # Remover arquivo anterior se existir
    output_file = 'documento_funcionalidades.docx'
    if os.path.exists(output_file):
        os.remove(output_file)

    # Criar novo documento
    doc = Document()

    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Aptos'
    font.size = Pt(11)

    # Título principal
    title = doc.add_heading('Documento de Funcionalidades - Sistema CONTEX', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adicionar seção 1
    doc.add_heading('1. Visão Geral do Sistema', 1)
    p = doc.add_paragraph()
    p.add_run('O sistema é uma aplicação web desenvolvida em PHP usando o framework CodeIgniter, destinada à gestão financeira pessoal. Permite aos usuários controlar suas finanças através de módulos específicos para lançamentos, faturas de cartão de crédito, investimentos, vendas, ordens de serviço e outros aspectos financeiros.')

    # Adicionar seção 2
    doc.add_heading('2. Arquitetura do Sistema', 1)

    doc.add_heading('2.1 Tecnologias Utilizadas', 2)
    tech_list = [
        'Framework: CodeIgniter 3.x',
        'Banco de Dados: MySQL/MariaDB',
        'Frontend: HTML, CSS, JavaScript, jQuery, Bootstrap',
        'Sessões: Database-driven sessions',
        'Autenticação: Baseada em permissões'
    ]
    for tech in tech_list:
        doc.add_paragraph(tech, style='List Bullet')

    doc.add_heading('2.2 Estrutura de Diretórios', 2)
    dir_list = [
        'application/: Código principal da aplicação',
        'system/: Framework CodeIgniter',
        'assets/: Recursos estáticos (CSS, JS, imagens)',
        'docs/: Documentação e scripts SQL'
    ]
    for dir_item in dir_list:
        doc.add_paragraph(dir_item, style='List Bullet')

    # Adicionar seção 3
    doc.add_heading('3. Módulos do Sistema', 1)

    # Módulo Financeiro
    doc.add_heading('3.1 Módulo Financeiro', 2)
    doc.add_paragraph('Responsabilidades: Controle completo das finanças pessoais')

    doc.add_heading('3.1.1 Submódulo Lançamentos', 3)
    lancamentos_list = [
        'Entradas: Registro de receitas (salários, rendimentos, etc.)',
        'Saídas: Registro de despesas (compras, contas, etc.)',
        'Filtros: Por período, status, tipo',
        'Características: Valores positivos para entradas, negativos para saídas'
    ]
    for item in lancamentos_list:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.1.2 Submódulo Faturas de Cartão', 3)
    faturas_list = [
        'Gestão de Cartões: Cadastro, edição, ativação/desativação',
        'Faturas: Criação automática por período',
        'Lançamentos: Compras à vista e parceladas',
        'Compras de Terceiros: Vinculação a clientes'
    ]
    for item in faturas_list:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.1.3 Submódulo Cartões', 3)
    cartoes_list = [
        'Cadastro de Cartões: Dados criptografados (número, validade, CVC)',
        'Cartões Titulares e Adicionais: Controle de titularidade',
        'Ativação/Desativação: Controle de status',
        'Associação Automática: Primeiro cartão vincula faturas existentes',
        'Exclusão Controlada: Validação de faturas associadas'
    ]
    for item in cartoes_list:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.1.4 Submódulo Investimentos', 3)
    investimentos_list = [
        'Aplicações: Registro de investimentos com débito automático',
        'Resgates: Retirada de investimentos com crédito automático',
        'Integração Financeira: Vinculação com módulo de lançamentos',
        'Filtros por Período: Controle temporal detalhado',
        'Status: Pendente/Efetivado'
    ]
    for item in investimentos_list:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.1.5 Submódulo Despesas', 3)
    despesas_list = [
        'Despesas Únicas e Recorrentes: Controle diferenciado',
        'Parcelamento Automático: Até 48 parcelas',
        'Integração com Lançamentos: Vinculação automática opcional',
        'Controle de Vencimento: Dia fixo por despesa',
        'Status: Pendente/Pago'
    ]
    for item in despesas_list:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.1.6 Submódulo Pendências', 3)
    pendencias_list = [
        'Contas a Receber/Pagar: Controle financeiro',
        'Vinculação com Clientes: Relacionamento direto',
        'Tipos: Crédito/Débito',
        'Pagamento: Registro de baixa com data',
        'Integração: Geração automática de lançamentos'
    ]
    for item in pendencias_list:
        doc.add_paragraph(item, style='List Bullet')

    # Módulo Clientes
    doc.add_heading('3.2 Módulo Clientes', 2)
    doc.add_paragraph('Responsabilidades: Gestão completa de clientes do sistema')

    doc.add_heading('Funcionalidades Principais:', 3)
    clientes_list = [
        'Cadastro de Clientes: Dados pessoais, contato, endereço',
        'Nome, CPF, telefone, email, data nascimento',
        'Endereço completo (logradouro, número, bairro, cidade, UF, CEP)',
        'Validação de dados obrigatórios',
        'Edição de Dados: Alteração completa das informações',
        'Visualização Detalhada: Perfil completo com histórico',
        'Exclusão Lógica: Desativação sem perda de dados',
        'Histórico Financeiro: Pendências de crédito e débito'
    ]
    for item in clientes_list:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Regras de Negócio:', 3)
    clientes_rules = [
        'Cliente vinculado ao usuário proprietário',
        'Dados padronizados (nomes, endereços)',
        'Controle de permissões por usuário',
        'Relacionamentos com OS, Vendas e Lançamentos'
    ]
    for item in clientes_rules:
        doc.add_paragraph(item, style='List Bullet')

    # Módulo Produtos
    doc.add_heading('3.3 Módulo Produtos', 2)
    doc.add_paragraph('Responsabilidades: Controle de estoque e produtos comercializados')

    doc.add_heading('Funcionalidades Principais:', 3)
    produtos_list = [
        'Cadastro de Produtos: Descrição, preços, controle de estoque',
        'Descrição, unidade, preço compra/venda',
        'Controle de estoque mínimo e atual',
        'Entradas e saídas automáticas',
        'Edição e Visualização: Gestão completa dos produtos',
        'Controle de Estoque: Atualização automática nas vendas/OS',
        'Exclusão Lógica: Manutenção da integridade referencial'
    ]
    for item in produtos_list:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Regras de Negócio:', 3)
    produtos_rules = [
        'Preços em formato monetário brasileiro',
        'Estoque atualizado automaticamente',
        'Validações de campos obrigatórios',
        'Controle de permissões de acesso'
    ]
    for item in produtos_rules:
        doc.add_paragraph(item, style='List Bullet')

    # Módulo Serviços
    doc.add_heading('3.4 Módulo Serviços', 2)
    doc.add_paragraph('Responsabilidades: Gestão de serviços oferecidos')

    doc.add_heading('Funcionalidades Principais:', 3)
    servicos_list = [
        'Cadastro de Serviços: Nome, descrição, preço',
        'Edição e Exclusão: Gestão completa dos serviços',
        'Integração com OS: Vinculação automática'
    ]
    for item in servicos_list:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Regras de Negócio:', 3)
    servicos_rules = [
        'Preços sem formatação (apenas números)',
        'Exclusão remove vínculos com OS',
        'Validações de campos obrigatórios'
    ]
    for item in servicos_rules:
        doc.add_paragraph(item, style='List Bullet')

    # Módulo OS
    doc.add_heading('3.5 Módulo Ordens de Serviço (OS)', 2)
    doc.add_paragraph('Responsabilidades: Controle de serviços técnicos e reparos')

    doc.add_heading('Funcionalidades Principais:', 3)
    os_list = [
        'Criação de OS: Cliente, responsável, descrição do problema',
        'Dados do equipamento, defeito declarado/encontrado',
        'Status: Aberto, Orçamento, Aprovado, Em Andamento, Cancelado, Pronto, Faturado',
        'Garantia, observações, laudo técnico',
        'Gestão de Produtos/Serviços: Adição/removal de itens',
        'Anexos: Upload de arquivos (fotos, documentos)',
        'Faturamento: Geração automática de lançamentos',
        'Impressão: Relatórios formatados'
    ]
    for item in os_list:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Regras de Negócio:', 3)
    os_rules = [
        'Controle de estoque automático nos produtos',
        'Faturamento gera lançamentos no financeiro',
        'Status controlado e obrigatório',
        'Anexos com tipos específicos permitidos',
        'Exclusão em cascata (produtos, serviços, anexos)'
    ]
    for item in os_rules:
        doc.add_paragraph(item, style='List Bullet')

    # Módulo Vendas
    doc.add_heading('3.6 Módulo Vendas', 2)
    doc.add_paragraph('Responsabilidades: Controle de vendas de produtos')

    doc.add_heading('Funcionalidades Principais:', 3)
    vendas_list = [
        'Criação de Vendas: Cliente, responsável, data',
        'Gestão de Produtos: Adição/removal de itens vendidos',
        'Controle de Estoque: Atualização automática',
        'Faturamento: Geração de lançamentos financeiros',
        'Impressão: Recibos e relatórios'
    ]
    for item in vendas_list:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Regras de Negócio:', 3)
    vendas_rules = [
        'Estoque atualizado automaticamente',
        'Faturamento gera receita no financeiro',
        'Controle de permissões granular',
        'Validações de campos obrigatórios'
    ]
    for item in vendas_rules:
        doc.add_paragraph(item, style='List Bullet')

    # Módulo Usuários
    doc.add_heading('3.7 Módulo Usuários e Permissões', 2)
    doc.add_paragraph('Responsabilidades: Controle de acesso e usuários do sistema')

    doc.add_heading('Funcionalidades Principais:', 3)
    usuarios_list = [
        'Gestão de Usuários: Cadastro, edição, ativação/desativação',
        'Perfis de Permissão: Controle granular de acessos',
        'Autenticação: Login seguro com sessões'
    ]
    for item in usuarios_list:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Regras de Negócio:', 3)
    usuarios_rules = [
        'Usuário administrador tem acesso total',
        'Permissões por módulo (ver, adicionar, editar, excluir)',
        'Controle de sessão obrigatório',
        'Senhas criptografadas'
    ]
    for item in usuarios_rules:
        doc.add_paragraph(item, style='List Bullet')

    # Módulo Relatórios
    doc.add_heading('3.8 Módulo Relatórios', 2)
    doc.add_paragraph('Responsabilidades: Geração de relatórios diversos')

    doc.add_heading('Funcionalidades Principais:', 3)
    relatorios_list = [
        'Relatórios Financeiros: Análise de receitas/despesas',
        'Relatórios de Clientes: Histórico e estatísticas',
        'Relatórios de Produtos/Serviços: Vendas e utilização',
        'Impressão: Formatada para diversos tipos'
    ]
    for item in relatorios_list:
        doc.add_paragraph(item, style='List Bullet')

    # Módulo Arquivos
    doc.add_heading('3.9 Módulo Arquivos/Documentos', 2)
    doc.add_paragraph('Responsabilidades: Gestão de documentos e arquivos')

    doc.add_heading('Funcionalidades Principais:', 3)
    arquivos_list = [
        'Upload e Organização: Arquivos por categoria',
        'Controle de Versões: Histórico de documentos',
        'Compartilhamento: Acesso controlado'
    ]
    for item in arquivos_list:
        doc.add_paragraph(item, style='List Bullet')

    # Módulo Conecte
    doc.add_heading('3.10 Módulo Conecte', 2)
    doc.add_paragraph('Responsabilidades: Integração e comunicação externa')

    doc.add_heading('Funcionalidades Principais:', 3)
    conecte_list = [
        'Integração Externa: APIs e sistemas terceiros',
        'Comunicação: Notificações e alertas',
        'Sincronização: Dados entre sistemas'
    ]
    for item in conecte_list:
        doc.add_paragraph(item, style='List Bullet')

    # Adicionar seção 4
    doc.add_heading('4. Regras de Negócio', 1)

    doc.add_heading('4.1 Autenticação e Autorização', 2)
    auth_list = [
        'Sistema baseado em sessões',
        'Controle de permissões por módulo',
        'Usuário administrador tem acesso total'
    ]
    for item in auth_list:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.2 Controle Financeiro', 2)
    financeiro_list = [
        'Saldo: Calculado dinamicamente (entradas - saídas)',
        'Períodos: Controle por mês/ano',
        'Status: Pendente/Efetivado',
        'Tipos: Entrada (tipo=1) / Saída (tipo=2)'
    ]
    for item in financeiro_list:
        doc.add_paragraph(item, style='List Bullet')

    # Adicionar seção 5
    doc.add_heading('5. Funcionalidades Principais', 1)

    doc.add_heading('5.1 Dashboard Financeiro', 2)
    dashboard_list = [
        'Saldo disponível',
        'Entradas/Saídas pendentes',
        'Totais por período',
        'Lançamentos ocultos'
    ]
    for item in dashboard_list:
        doc.add_paragraph(item, style='List Bullet')

    # Adicionar seção 6
    doc.add_heading('6. Regras Técnicas', 1)

    doc.add_heading('6.1 Versionamento', 2)
    doc.add_paragraph('Formato: YYYY.Q.R (Ano.Trimestre.Release)')

    doc.add_heading('6.2 Commits', 2)
    commits_list = [
        'Mensagens curtas em inglês',
        'Solicitar permissão antes de push para origin'
    ]
    for item in commits_list:
        doc.add_paragraph(item, style='List Bullet')

    # Salvar documento
    doc.save('documento_funcionalidades.docx')
    print('Documento Word criado com sucesso!')

if __name__ == '__main__':
    create_word_document()