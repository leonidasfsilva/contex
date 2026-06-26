import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_requirements_document():
    # Verificar se arquivo já existe e remover
    import os
    doc_path = 'docs/documento_requisitos.docx'
    if os.path.exists(doc_path):
        os.remove(doc_path)
        print(f"Arquivo existente '{doc_path}' removido.")

    # Criar documento
    doc = Document()

    # Configurar fonte Aptos para todo o documento
    from docx.shared import Pt
    from docx.oxml.ns import nsdecls, qn
    from docx.oxml import parse_xml

    # Aplicar fonte Aptos a todos os estilos existentes
    for style in doc.styles:
        if hasattr(style, 'font'):
            style.font.name = 'Aptos'
            style.font.size = Pt(11)

    # Configurar fonte Aptos para estilos de título via XML
    for i in range(1, 10):  # Heading 1 até Heading 9
        style_name = f'Heading {i}'
        if style_name in doc.styles:
            heading_style = doc.styles[style_name]
            rPr = heading_style.element.get_or_add_rPr()
            rFonts = parse_xml(r'<w:rFonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:ascii="Aptos" w:hAnsi="Aptos" w:cs="Aptos"/>')
            rPr.append(rFonts)

    # Título principal
    title = doc.add_heading('Documento de Requisitos - Sistema CONTEX', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Configurar fonte Aptos e tamanho 20pt para o título principal
    for run in title.runs:
        run.font.name = 'Aptos'
        run.font.size = Pt(20)

    # Introdução
    doc.add_heading('1. Introdução', level=1)
    doc.add_paragraph(
        'Este documento descreve os requisitos funcionais e não funcionais do Sistema de Gestão Financeira Pessoal. '
        'O sistema é desenvolvido em PHP usando o framework CodeIgniter '
        'e MySQL como banco de dados.'
    )

    # Visão Geral do Sistema
    doc.add_heading('2. Visão Geral do Sistema', level=1)
    doc.add_paragraph(
        'O sistema é uma aplicação web para gestão financeira pessoal que permite aos usuários controlar suas '
        'finanças através de lançamentos de receitas e despesas, gerenciamento de cartões de crédito, faturas, '
        'clientes, produtos, serviços e ordens de serviço.'
    )

    # Arquitetura do Sistema
    doc.add_heading('3. Arquitetura do Sistema', level=1)

    # Tecnologias
    doc.add_heading('3.1 Tecnologias Utilizadas', level=2)
    tech_table = doc.add_table(rows=1, cols=2)
    tech_table.style = 'Table Grid'

    # Cabeçalho da tabela
    header_cells = tech_table.rows[0].cells
    header_cells[0].text = 'Componente'
    header_cells[1].text = 'Tecnologia'

    # Dados da tabela
    technologies = [
        ('Framework', 'CodeIgniter 3.x'),
        ('Linguagem', 'PHP 7.x+'),
        ('Banco de Dados', 'MySQL/MariaDB'),
        ('Frontend', 'HTML5, CSS3, JavaScript, jQuery'),
        ('UI Framework', 'Bootstrap, Font Awesome'),
        ('Charts', 'Chart.js'),
        ('PDF Generation', 'MPDF'),
        ('Session Management', 'Database Sessions'),
        ('Authentication', 'Custom Authentication'),
        ('File Upload', 'CodeIgniter Upload Library'),
        ('Validation', 'jQuery Validate, CodeIgniter Form Validation')
    ]

    for tech, version in technologies:
        row_cells = tech_table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = version

    # Estrutura de Diretórios
    doc.add_heading('3.2 Estrutura de Diretórios', level=2)
    doc.add_paragraph(
        'O sistema segue a estrutura padrão do CodeIgniter com diretórios específicos para cada módulo:'
    )

    dir_structure = """
application/
├── config/          # Arquivos de configuração
├── controllers/     # Controladores (Clientes, Produtos, OS, Financeiro)
├── models/          # Modelos de dados
├── views/           # Templates das páginas
├── helpers/         # Funções auxiliares
└── libraries/       # Bibliotecas customizadas

assets/              # Arquivos estáticos (CSS, JS, imagens)
system/              # Core do CodeIgniter
docs/                # Documentação e scripts SQL
"""
    doc.add_paragraph(dir_structure)

    # Módulos do Sistema
    doc.add_heading('4. Módulos do Sistema', level=1)

    # 4.1 Módulo de Usuários
    doc.add_heading('4.1 Módulo de Usuários', level=2)
    doc.add_paragraph(
        'Gerencia usuários do sistema com diferentes níveis de permissão e controle de acesso.'
    )

    # Usuários - Requisitos Funcionais Detalhados
    doc.add_heading('4.1.1 Cadastro de Usuários', level=3)
    user_cad_reqs = doc.add_paragraph()
    user_cad_reqs.add_run('RF-001:').bold = True
    user_cad_reqs.add_run(' Sistema deve permitir cadastro de usuários com os campos: nome, RG, CPF (único), email (único), telefone, celular, senha, rua, número, bairro, cidade, estado, CEP\n')
    user_cad_reqs.add_run('RF-002:').bold = True
    user_cad_reqs.add_run(' Validação obrigatória de CPF válido e único no sistema\n')
    user_cad_reqs.add_run('RF-003:').bold = True
    user_cad_reqs.add_run(' Validação obrigatória de email válido e único no sistema\n')
    user_cad_reqs.add_run('RF-004:').bold = True
    user_cad_reqs.add_run(' Senha deve ter mínimo 6 caracteres e ser criptografada com hash seguro\n')
    user_cad_reqs.add_run('RF-005:').bold = True
    user_cad_reqs.add_run(' Usuário deve ser vinculado obrigatoriamente a uma permissão/perfil\n')
    user_cad_reqs.add_run('RF-006:').bold = True
    user_cad_reqs.add_run(' Status inicial do usuário deve ser ativo (situacao = 1)\n')
    user_cad_reqs.add_run('RF-007:').bold = True
    user_cad_reqs.add_run(' Data de cadastro deve ser registrada automaticamente')

    doc.add_heading('4.1.2 Autenticação e Segurança', level=3)
    user_auth_reqs = doc.add_paragraph()
    user_auth_reqs.add_run('RF-008:').bold = True
    user_auth_reqs.add_run(' Sistema deve validar login com email/CPF e senha\n')
    user_auth_reqs.add_run('RF-009:').bold = True
    user_auth_reqs.add_run(' Sessão deve expirar após 2 horas de inatividade\n')
    user_auth_reqs.add_run('RF-010:').bold = True
    user_auth_reqs.add_run(' Bloqueio automático após 3 tentativas de login fracassadas\n')
    user_auth_reqs.add_run('RF-011:').bold = True
    user_auth_reqs.add_run(' Controle de acesso baseado em permissões por módulo\n')
    user_auth_reqs.add_run('RF-012:').bold = True
    user_auth_reqs.add_run(' Usuário administrador pode gerenciar todos os módulos\n')
    user_auth_reqs.add_run('RF-013:').bold = True
    user_auth_reqs.add_run(' Logout deve destruir completamente a sessão')

    doc.add_heading('4.1.3 Controle de Permissões', level=3)
    user_perm_reqs = doc.add_paragraph()
    user_perm_reqs.add_run('RF-014:').bold = True
    user_perm_reqs.add_run(' Sistema deve ter perfis pré-definidos (Administrador) com permissões específicas\n')
    user_perm_reqs.add_run('RF-015:').bold = True
    user_perm_reqs.add_run(' Cada módulo deve ter permissões granulares: visualizar (v), adicionar (a), editar (e), excluir (d)\n')
    user_perm_reqs.add_run('RF-016:').bold = True
    user_perm_reqs.add_run(' Permissões devem ser armazenadas como array serializado no banco\n')
    user_perm_reqs.add_run('RF-017:').bold = True
    user_perm_reqs.add_run(' Verificação de permissões deve ocorrer em cada action do controller\n')
    user_perm_reqs.add_run('RF-018:').bold = True
    user_perm_reqs.add_run(' Usuários sem permissão devem ser redirecionados com mensagem de erro')

    # 4.2 Módulo Financeiro
    doc.add_heading('4.2 Módulo Financeiro', level=2)
    doc.add_paragraph(
        'Núcleo do sistema responsável pelo controle financeiro pessoal.'
    )

    # Lançamentos
    doc.add_heading('4.2.1 Lançamentos Financeiros', level=3)
    lanc_reqs = doc.add_paragraph()
    lanc_reqs.add_run('RF-019:').bold = True
    lanc_reqs.add_run(' Sistema deve permitir registrar lançamentos com campos obrigatórios: descricao (varchar 255, não vazio), valor (decimal 10,2, != 0), data_lancamento (date), tipo (tinyint 1=entrada/2=saída), id_usuario (int, preenchido automaticamente)\n')
    lanc_reqs.add_run('RF-020:').bold = True
    lanc_reqs.add_run(' Campos opcionais: data_pagamento (date), cliente_fornecedor (varchar 255), forma_pgto (int), baixado (tinyint, default 0), oculto (tinyint, default 0), observacoes (text)\n')
    lanc_reqs.add_run('RF-021:').bold = True
    lanc_reqs.add_run(' Validação obrigatória: descricao.trim() não vazio, valor != 0, data_lancamento válida, tipo em (1,2), id_usuario = session user ID\n')
    lanc_reqs.add_run('RF-022:').bold = True
    lanc_reqs.add_run(' Validação condicional: se baixado=1 então data_pagamento deve ser preenchida; se data_pagamento informada então deve ser >= data_lancamento\n')
    lanc_reqs.add_run('RF-023:').bold = True
    lanc_reqs.add_run(' Validação de formato: valor deve aceitar apenas números decimais válidos (ex: 123.45 ou 123,45); data deve estar no formato DD/MM/YYYY\n')
    lanc_reqs.add_run('RF-024:').bold = True
    lanc_reqs.add_run(' Campo baixado deve controlar status: 0=pendente (não afeta saldo), 1=pago (afeta saldo); campo oculto: 0=visível, 1=oculto da visualização padrão\n')
    lanc_reqs.add_run('RF-025:').bold = True
    lanc_reqs.add_run(' Sistema deve calcular saldo total: SELECT SUM(CASE WHEN baixado=1 THEN valor ELSE 0 END) FROM lancamentos WHERE id_usuario = ? AND oculto = 0\n')
    lanc_reqs.add_run('RF-026:').bold = True
    lanc_reqs.add_run(' Sistema deve calcular saldos parciais: entradas (SUM valor WHERE tipo=1 AND baixado=1), saidas (SUM valor WHERE tipo=2 AND baixado=1), pendentes (SUM valor WHERE baixado=0)\n')
    lanc_reqs.add_run('RF-027:').bold = True
    lanc_reqs.add_run(' Filtros devem incluir: periodo (3dias/5dias/7dias/15dias/30dias/60dias/90dias/especifico/mensal), status (pendente/efetivado/todos), tipo (entrada/saida/todos), dataInicial/dataFinal para periodo especifico\n')
    lanc_reqs.add_run('RF-028:').bold = True
    lanc_reqs.add_run(' Pesquisa deve funcionar por: descricao LIKE "%%term%%" OR cliente_fornecedor LIKE "%%term%%" OR observacoes LIKE "%%term%%"\n')
    lanc_reqs.add_run('RF-029:').bold = True
    lanc_reqs.add_run(' Sistema deve exibir lista paginada (30 registros por página) ordenada por data_lancamento DESC, id_lancamento DESC\n')
    lanc_reqs.add_run('RF-030:').bold = True
    lanc_reqs.add_run(' Auto-complete deve funcionar para campos descricao e cliente_fornecedor, retornando valores únicos dos últimos 100 lançamentos do usuário\n')
    lanc_reqs.add_run('RF-031:').bold = True
    lanc_reqs.add_run(' Exclusão deve ser lógica: UPDATE lancamentos SET status = 0 WHERE id_lancamento = ? AND id_usuario = ?\n')
    lanc_reqs.add_run('RF-032:').bold = True
    lanc_reqs.add_run(' Validação de edição: usuário só pode editar seus próprios lançamentos; lançamentos já baixados podem ter data_pagamento alterada mas não valor\n')
    lanc_reqs.add_run('RF-033:').bold = True
    lanc_reqs.add_run(' Sistema deve permitir cópia de lançamentos: duplicar registro alterando apenas data_lancamento\n')
    lanc_reqs.add_run('RF-034:').bold = True
    lanc_reqs.add_run(' Campo cliente_fornecedor deve ser usado para agrupar lançamentos similares e gerar relatórios por fornecedor/cliente')

    # Cartões de Crédito
    doc.add_heading('4.2.2 Cartões de Crédito', level=3)
    cartao_reqs = doc.add_paragraph()
    cartao_reqs.add_run('RF-035:').bold = True
    cartao_reqs.add_run(' Sistema deve permitir cadastro de cartões com campos obrigatórios: bandeira (varchar 45), numero (varchar 255, criptografado), dia_vencimento (int 1-31), id_usuario (int, preenchido automaticamente)\n')
    cartao_reqs.add_run('RF-036:').bold = True
    cartao_reqs.add_run(' Campos opcionais: validade (varchar 10), limite (decimal 10,2), apelido (varchar 45), status (tinyint, default 1), id_usuario_titular (int), adicional (tinyint, default 0)\n')
    cartao_reqs.add_run('RF-037:').bold = True
    cartao_reqs.add_run(' Validação obrigatória: bandeira não vazia, numero com 13-19 dígitos, dia_vencimento entre 1-31, id_usuario = session user ID\n')
    cartao_reqs.add_run('RF-038:').bold = True
    cartao_reqs.add_run(' Criptografia: numero deve ser criptografado usando encriptar() antes de salvar no banco\n')
    cartao_reqs.add_run('RF-039:').bold = True
    cartao_reqs.add_run(' Validação de formato: numero deve aceitar apenas números, mínimo 13 máximo 19 dígitos\n')
    cartao_reqs.add_run('RF-040:').bold = True
    cartao_reqs.add_run(' Sistema deve identificar cartão principal: SELECT * FROM cartoes WHERE id_usuario = ? AND status = 1 ORDER BY id_cartao LIMIT 1\n')
    cartao_reqs.add_run('RF-041:').bold = True
    cartao_reqs.add_run(' Cartões adicionais: adicional = 1, id_usuario_titular = id_cartao do titular, id_usuario = id do usuário adicional\n')
    cartao_reqs.add_run('RF-042:').bold = True
    cartao_reqs.add_run(' Validação de permissão: SELECT * FROM cartoes WHERE id_cartao = ? AND (id_usuario = ? OR id_usuario_titular = ?)\n')
    cartao_reqs.add_run('RF-043:').bold = True
    cartao_reqs.add_run(' Configuração de vencimento deve propagar: UPDATE faturas SET vencimento = CONCAT(ano_referencia, "-", mes_referencia, "-", ?) WHERE id_cartao = ? AND fatura_aberta IN (1,2)\n')
    cartao_reqs.add_run('RF-044:').bold = True
    cartao_reqs.add_run(' Sistema deve permitir busca por: bandeira LIKE "%%term%%" OR apelido LIKE "%%term%%"\n')
    cartao_reqs.add_run('RF-045:').bold = True
    cartao_reqs.add_run(' Sistema deve exibir lista paginada (20 por página) ordenada por status DESC, bandeira ASC\n')
    cartao_reqs.add_run('RF-046:').bold = True
    cartao_reqs.add_run(' Validação de edição: usuário titular pode editar cartões próprios e adicionais; usuário adicional só pode ver, não editar\n')
    cartao_reqs.add_run('RF-047:').bold = True
    cartao_reqs.add_run(' Exclusão deve ser lógica: UPDATE cartoes SET status = 0 WHERE id_cartao = ? (não excluir fisicamente por integridade referencial)')

    # Faturas
    doc.add_heading('4.2.3 Faturas', level=3)
    fatura_reqs = doc.add_paragraph()
    fatura_reqs.add_run('RF-048:').bold = True
    fatura_reqs.add_run(' Sistema deve criar faturas automaticamente: INSERT INTO faturas (id_usuario, id_cartao, mes_referencia, ano_referencia, vencimento, fatura_aberta) VALUES (?, ?, ?, ?, CONCAT(?, "-", ?, "-", dia_vencimento), 2)\n')
    fatura_reqs.add_run('RF-049:').bold = True
    fatura_reqs.add_run(' Campos obrigatórios: id_usuario (int), id_cartao (int, FK), mes_referencia (int 1-12), ano_referencia (int), vencimento (date), fatura_aberta (tinyint, default 2)\n')
    fatura_reqs.add_run('RF-050:').bold = True
    fatura_reqs.add_run(' Campos opcionais: fatura_paga (tinyint, default 0), data_pagamento (date), forma_pgto (int), valor_total (decimal 10,2), fatura_vinculada (tinyint, default 0), status (tinyint, default 1)\n')
    fatura_reqs.add_run('RF-051:').bold = True
    fatura_reqs.add_run(' Validação obrigatória: id_cartao existe na tabela cartoes, mes_referencia 1-12, ano_referencia >= ano atual, vencimento válida\n')
    fatura_reqs.add_run('RF-052:').bold = True
    fatura_reqs.add_run(' Compras parceladas: INSERT INTO lancamentos_faturas_assoc (id_lancamento, id_fatura, valor_parcela, data_compra, n_parcela, total_parcelas) para cada mês\n')
    fatura_reqs.add_run('RF-053:').bold = True
    fatura_reqs.add_run(' Compras à vista: INSERT INTO lancamentos_faturas_assoc (id_lancamento, id_fatura, valor_parcela, data_compra, n_parcela=1, total_parcelas=1)\n')
    fatura_reqs.add_run('RF-054:').bold = True
    fatura_reqs.add_run(' Cálculo valor_total: UPDATE faturas SET valor_total = (SELECT SUM(valor_parcela) FROM lancamentos_faturas_assoc WHERE id_fatura = faturas.id_fatura) WHERE id_fatura = ?\n')
    fatura_reqs.add_run('RF-055:').bold = True
    fatura_reqs.add_run(' Fechar fatura: UPDATE faturas SET fatura_aberta = 0, fatura_paga = 2 WHERE id_fatura = ?; criar próxima fatura automaticamente\n')
    fatura_reqs.add_run('RF-056:').bold = True
    fatura_reqs.add_run(' Pagar fatura: UPDATE faturas SET fatura_paga = 1, data_pagamento = ?, forma_pgto = ? WHERE id_fatura = ?\n')
    fatura_reqs.add_run('RF-057:').bold = True
    fatura_reqs.add_run(' Compras de terceiros: INSERT INTO pendencias (id_lancamento_fatura, id_usuario, id_cliente, descricao, tipo=1, valor, data_vencimento)\n')
    fatura_reqs.add_run('RF-058:').bold = True
    fatura_reqs.add_run(' Validação: faturas fechadas (fatura_aberta=0) não permitem novos lançamentos\n')
    fatura_reqs.add_run('RF-059:').bold = True
    fatura_reqs.add_run(' Saldos calculados: vencidas (vencimento < CURDATE() AND fatura_paga = 0), pendentes (fatura_paga = 0), pagas (fatura_paga = 1)\n')
    fatura_reqs.add_run('RF-060:').bold = True
    fatura_reqs.add_run(' Vinculação financeira: UPDATE faturas SET fatura_vinculada = 1 WHERE id_fatura = ?; INSERT INTO lancamentos (id_fatura, valor= -valor_total, tipo=2, baixado=1, ...)\n')
    fatura_reqs.add_run('RF-061:').bold = True
    fatura_reqs.add_run(' Sistema deve permitir reabrir fatura fechada: UPDATE faturas SET fatura_aberta = 1 WHERE id_fatura = ? AND fatura_paga = 2\n')
    fatura_reqs.add_run('RF-062:').bold = True
    fatura_reqs.add_run(' Filtros devem incluir: id_cartao (select), mes_referencia/ano_referencia, fatura_aberta, fatura_paga\n')
    fatura_reqs.add_run('RF-063:').bold = True
    fatura_reqs.add_run(' Ordenação padrão: ano_referencia DESC, mes_referencia DESC\n')
    fatura_reqs.add_run('RF-064:').bold = True
    fatura_reqs.add_run(' Validação de permissão: usuário só pode ver/editar faturas de seus cartões (próprios ou adicionais)')

    # Despesas
    doc.add_heading('4.2.4 Despesas', level=3)
    despesa_reqs = doc.add_paragraph()
    despesa_reqs.add_run('RF-065:').bold = True
    despesa_reqs.add_run(' Sistema deve permitir cadastro de despesas com campos obrigatórios: descricao (varchar 255, não vazio), valor_parcela (decimal 10,2 > 0), dia_vencimento (int 1-31), id_usuario (int, preenchido automaticamente)\n')
    despesa_reqs.add_run('RF-066:').bold = True
    despesa_reqs.add_run(' Campos opcionais: fornecedor (varchar 255), nome_terceiro (varchar 255), valor_total (decimal 10,2), total_parcelas (int), despesa_parcelada (tinyint, default 0), despesa_terceiros (tinyint, default 0), despesa_oculta (tinyint, default 0), id_forma_pagamento (int)\n')
    despesa_reqs.add_run('RF-067:').bold = True
    despesa_reqs.add_run(' Validação obrigatória: descricao.trim() não vazio, valor_parcela > 0, dia_vencimento 1-31, id_usuario = session user ID\n')
    despesa_reqs.add_run('RF-068:').bold = True
    despesa_reqs.add_run(' Despesas recorrentes: despesa_parcelada = 0, sistema cria registros mensais automaticamente baseado em dia_vencimento\n')
    despesa_reqs.add_run('RF-069:').bold = True
    despesa_reqs.add_run(' Despesas parceladas: despesa_parcelada = 1, total_parcelas > 1, valor_total = valor_parcela * total_parcelas\n')
    despesa_reqs.add_run('RF-070:').bold = True
    despesa_reqs.add_run(' Cada parcela gera registro: INSERT INTO registros_despesas (id_despesa, data_vencimento, valor, registro_pago) VALUES (?, DATE(CONCAT(ano, "-", mes, "-", dia_vencimento)), valor_parcela, 0)\n')
    despesa_reqs.add_run('RF-071:').bold = True
    despesa_reqs.add_run(' Integração financeira automática: INSERT INTO lancamentos (descricao, valor= -valor, data_lancamento= data_vencimento, tipo=2, baixado= registro_pago, id_usuario) para cada parcela\n')
    despesa_reqs.add_run('RF-072:').bold = True
    despesa_reqs.add_run(' Despesas de terceiros: despesa_terceiros = 1, INSERT INTO pendencias (id_cliente, descricao, tipo=2, valor= -valor, data_vencimento)\n')
    despesa_reqs.add_run('RF-073:').bold = True
    despesa_reqs.add_run(' Filtros devem incluir: periodo (3dias/5dias/7dias/15dias/30dias/60dias/90dias/mensal), status (pendente/pago/todos), tipo (unica/recorrente/parcelada/terceiros)\n')
    despesa_reqs.add_run('RF-074:').bold = True
    despesa_reqs.add_run(' Pagamento: UPDATE registros_despesas SET registro_pago = 1, data_pagamento = ? WHERE id_registro = ?; UPDATE lancamentos SET baixado = 1, data_pagamento = ? WHERE id_despesa = ?\n')
    despesa_reqs.add_run('RF-075:').bold = True
    despesa_reqs.add_run(' Sistema deve permitir configuração global: ativar/desativar integração automática com financeiro\n')
    despesa_reqs.add_run('RF-076:').bold = True
    despesa_reqs.add_run(' Exclusão lógica: UPDATE despesas SET status = 0 WHERE id_despesa = ?; UPDATE registros_despesas SET status = 0 WHERE id_despesa = ?\n')
    despesa_reqs.add_run('RF-077:').bold = True
    despesa_reqs.add_run(' Campo despesa_oculta: 0=visível no financeiro, 1=oculta (despesas pessoais não compartilhadas)\n')
    despesa_reqs.add_run('RF-078:').bold = True
    despesa_reqs.add_run(' Validação: dia_vencimento deve ser válido para o mês (ex: dia 31 não válido para fevereiro)\n')
    despesa_reqs.add_run('RF-079:').bold = True
    despesa_reqs.add_run(' Sistema deve calcular totais: total_despesas (SUM valor WHERE registro_pago=1), despesas_pendentes (SUM valor WHERE registro_pago=0)')

    # Pendências
    doc.add_heading('4.2.5 Pendências', level=3)
    pendencia_reqs = doc.add_paragraph()
    pendencia_reqs.add_run('RF-080:').bold = True
    pendencia_reqs.add_run(' Sistema deve permitir cadastro de pendências com campos obrigatórios: id_cliente (int, FK válido), descricao (varchar 255, não vazio), tipo (tinyint 1=crédito/2=débito), valor (decimal 10,2 != 0), data_vencimento (date), id_usuario (int, preenchido automaticamente)\n')
    pendencia_reqs.add_run('RF-081:').bold = True
    pendencia_reqs.add_run(' Campos opcionais: data_pagamento (date), quitado (tinyint, default 0), forma_pagamento (int), id_lancamento_fatura (int), status (tinyint, default 1)\n')
    pendencia_reqs.add_run('RF-082:').bold = True
    pendencia_reqs.add_run(' Validação obrigatória: id_cliente existe na tabela clientes, descricao.trim() não vazio, tipo em (1,2), valor != 0, data_vencimento >= CURDATE(), id_usuario = session user ID\n')
    pendencia_reqs.add_run('RF-083:').bold = True
    pendencia_reqs.add_run(' Valores de débito (tipo=2) devem ser armazenados como positivos no banco (representam valores a receber); créditos (tipo=1) são valores a pagar\n')
    pendencia_reqs.add_run('RF-084:').bold = True
    pendencia_reqs.add_run(' Cálculo totais por cliente: creditos (SUM valor WHERE tipo=1 AND quitado=0), debitos (SUM valor WHERE tipo=2 AND quitado=0), saldo (debitos - creditos)\n')
    pendencia_reqs.add_run('RF-085:').bold = True
    pendencia_reqs.add_run(' Cálculo totais gerais: SELECT SUM(CASE WHEN tipo=1 THEN valor ELSE -valor END) FROM pendencias WHERE id_usuario = ? AND quitado = 0\n')
    pendencia_reqs.add_run('RF-086:').bold = True
    pendencia_reqs.add_run(' Quitação: UPDATE pendencias SET quitado = 1, data_pagamento = ? WHERE id_pendencia = ?; INSERT INTO lancamentos (valor= CASE WHEN tipo=1 THEN -valor ELSE valor END, tipo= CASE WHEN tipo=1 THEN 2 ELSE 1 END, baixado=1, ...)\n')
    pendencia_reqs.add_run('RF-087:').bold = True
    pendencia_reqs.add_run(' Sistema deve permitir gerar lançamento financeiro automático ao quitar: checkbox "gerar_lancamento" cria registro em lancamentos\n')
    pendencia_reqs.add_run('RF-088:').bold = True
    pendencia_reqs.add_run(' Filtros devem incluir: id_cliente (select), data_vencimento BETWEEN periodo, quitado (0/1/todos), tipo (credito/debito/todos)\n')
    pendencia_reqs.add_run('RF-089:').bold = True
    pendencia_reqs.add_run(' Ordenação padrão: data_vencimento ASC (vencimentos mais próximos primeiro)\n')
    pendencia_reqs.add_run('RF-090:').bold = True
    pendencia_reqs.add_run(' Validação de edição: usuário só pode editar suas próprias pendências; pendências quitadas não podem ser editadas\n')
    pendencia_reqs.add_run('RF-091:').bold = True
    pendencia_reqs.add_run(' Exclusão lógica: UPDATE pendencias SET status = 0 WHERE id_pendencia = ? AND quitado = 0\n')
    pendencia_reqs.add_run('RF-092:').bold = True
    pendencia_reqs.add_run(' Campo id_lancamento_fatura vincula pendências criadas automaticamente por compras de terceiros em faturas\n')
    pendencia_reqs.add_run('RF-093:').bold = True
    pendencia_reqs.add_run(' Sistema deve exibir alertas visuais: pendências vencidas (data_vencimento < CURDATE() AND quitado = 0), vencendo hoje, vencendo amanhã\n')
    pendencia_reqs.add_run('RF-094:').bold = True
    pendencia_reqs.add_run(' Relatório deve incluir colunas calculadas: dias_atraso (DATEDIFF(CURDATE(), data_vencimento)), status (IF(quitado=1,"Quitada",IF(data_vencimento<CURDATE(),"Vencida","Pendente")))')

    # Investimentos
    doc.add_heading('4.2.6 Investimentos', level=3)
    investimento_reqs = doc.add_paragraph()
    investimento_reqs.add_run('RF-095:').bold = True
    investimento_reqs.add_run(' Sistema deve permitir registrar aplicações/resgates com campos obrigatórios: descricao (varchar 255, não vazio), valor (decimal 10,2 != 0), data_lancamento (date), tipo (tinyint 1=aplicação/2=resgate), id_usuario (int, preenchido automaticamente)\n')
    investimento_reqs.add_run('RF-096:').bold = True
    investimento_reqs.add_run(' Campos opcionais: forma_pgto (int), baixado (tinyint, default 0), debito_conta (tinyint, default 0), status (tinyint, default 1)\n')
    investimento_reqs.add_run('RF-097:').bold = True
    investimento_reqs.add_run(' Validação obrigatória: descricao.trim() não vazio, valor != 0, data_lancamento válida, tipo em (1,2), id_usuario = session user ID\n')
    investimento_reqs.add_run('RF-098:').bold = True
    investimento_reqs.add_run(' Aplicações (tipo=1) devem ter valores positivos; resgates (tipo=2) valores positivos (representam saída de investimento)\n')
    investimento_reqs.add_run('RF-099:').bold = True
    investimento_reqs.add_run(' Cálculo saldos: entradas (SUM valor WHERE tipo=1 AND baixado=1), saidas (SUM valor WHERE tipo=2 AND baixado=1), total_disponivel (entradas - saidas)\n')
    investimento_reqs.add_run('RF-100:').bold = True
    investimento_reqs.add_run(' Integração financeira automática: IF debito_conta=1 THEN INSERT INTO lancamentos (valor= CASE WHEN tipo=1 THEN -valor ELSE valor END, tipo= CASE WHEN tipo=1 THEN 2 ELSE 1 END, baixado=1, descricao= "Investimento: " + descricao, ...)\n')
    investimento_reqs.add_run('RF-101:').bold = True
    investimento_reqs.add_run(' Checkbox debito_conta: quando marcado, cria lançamento financeiro automático debitado/creditado da conta corrente\n')
    investimento_reqs.add_run('RF-102:').bold = True
    investimento_reqs.add_run(' Filtros devem incluir: periodo (3dias/5dias/7dias/15dias/30dias/60dias/90dias), status (pendente/efetivado/todos), tipo (aplicacao/resgate/todos)\n')
    investimento_reqs.add_run('RF-103:').bold = True
    investimento_reqs.add_run(' Ordenação padrão: data_lancamento DESC, id_investimento DESC\n')
    investimento_reqs.add_run('RF-104:').bold = True
    investimento_reqs.add_run(' Paginação deve exibir 20 registros por página\n')
    investimento_reqs.add_run('RF-105:').bold = True
    investimento_reqs.add_run(' Exclusão lógica: UPDATE investimentos SET status = 0 WHERE id_investimento = ? AND baixado = 0\n')
    investimento_reqs.add_run('RF-106:').bold = True
    investimento_reqs.add_run(' Validação: investimentos já baixados não podem ser editados/excluídos\n')
    investimento_reqs.add_run('RF-107:').bold = True
    investimento_reqs.add_run(' Campo baixado: 0=pendente (não efetivado), 1=efetivado (já ocorreu a aplicação/resgate)\n')
    investimento_reqs.add_run('RF-108:').bold = True
    investimento_reqs.add_run(' Sistema deve permitir configuração de tipos de investimento (ações, fundos, CDB, etc.) através de tabela relacionada\n')
    investimento_reqs.add_run('RF-109:').bold = True
    investimento_reqs.add_run(' Relatório deve incluir rentabilidade: ((total_disponivel - total_aplicado) / total_aplicado) * 100')

    # 4.3 Módulo de Clientes
    doc.add_heading('4.3 Módulo de Clientes', level=2)
    doc.add_paragraph('Gerencia cadastro e manutenção de clientes do sistema.')

    doc.add_heading('4.3.1 Cadastro de Clientes', level=3)
    cliente_cad_reqs = doc.add_paragraph()
    cliente_cad_reqs.add_run('RF-019:').bold = True
    cliente_cad_reqs.add_run(' Sistema deve permitir cadastro de clientes com campos obrigatórios: nomeCliente (varchar 255, não vazio), pessoa_fisica (tinyint 1=PF/0=PJ), documento (varchar 20, único), telefone (varchar 20, não vazio), email (varchar 100, válido e único)\n')
    cliente_cad_reqs.add_run('RF-020:').bold = True
    cliente_cad_reqs.add_run(' Campos opcionais: sexo (varchar 20), celular (varchar 20), rua (varchar 70), numero (varchar 15), bairro (varchar 45), cidade (varchar 45), estado (varchar 20), cep (varchar 20)\n')
    cliente_cad_reqs.add_run('RF-021:').bold = True
    cliente_cad_reqs.add_run(' Validação obrigatória: nomeCliente.trim() não vazio, documento único na tabela clientes, email válido (formato) e único na tabela clientes\n')
    cliente_cad_reqs.add_run('RF-022:').bold = True
    cliente_cad_reqs.add_run(' Validação condicional: se pessoa_fisica=1 então documento deve ter 11 dígitos + validação CPF; se pessoa_fisica=0 então documento deve ter 14 dígitos + validação CNPJ\n')
    cliente_cad_reqs.add_run('RF-023:').bold = True
    cliente_cad_reqs.add_run(' Campo dataCadastro deve ser preenchido automaticamente com CURRENT_TIMESTAMP na criação\n')
    cliente_cad_reqs.add_run('RF-024:').bold = True
    cliente_cad_reqs.add_run(' Sistema deve permitir edição completa de todos os campos exceto dataCadastro\n')
    cliente_cad_reqs.add_run('RF-025:').bold = True
    cliente_cad_reqs.add_run(' Exclusão deve ser lógica através de campo situacao=0 (não implementado na tabela atual, deve ser adicionado)\n')
    cliente_cad_reqs.add_run('RF-026:').bold = True
    cliente_cad_reqs.add_run(' Validação de telefone: deve aceitar apenas números, mínimo 10 dígitos (DDD + número)\n')
    cliente_cad_reqs.add_run('RF-027:').bold = True
    cliente_cad_reqs.add_run(' Validação de CEP: deve aceitar apenas números, exatamente 8 dígitos quando informado')

    doc.add_heading('4.3.2 Vinculações e Consultas', level=3)
    cliente_vinc_reqs = doc.add_paragraph()
    cliente_vinc_reqs.add_run('RF-028:').bold = True
    cliente_vinc_reqs.add_run(' Sistema deve permitir vincular clientes a lançamentos financeiros através do campo cliente_fornecedor (varchar 255, pode ser nome do cliente)\n')
    cliente_vinc_reqs.add_run('RF-029:').bold = True
    cliente_vinc_reqs.add_run(' Sistema deve permitir vincular clientes a ordens de serviço através do campo clientes_id (int, FK para clientes.idClientes)\n')
    cliente_vinc_reqs.add_run('RF-030:').bold = True
    cliente_vinc_reqs.add_run(' Sistema deve permitir vincular clientes a vendas através do campo clientes_id (int, FK para clientes.idClientes)\n')
    cliente_vinc_reqs.add_run('RF-031:').bold = True
    cliente_vinc_reqs.add_run(' Sistema deve permitir vincular clientes a equipamentos através do campo clientes_id (int, FK para clientes.idClientes)\n')
    cliente_vinc_reqs.add_run('RF-032:').bold = True
    cliente_vinc_reqs.add_run(' Sistema deve permitir busca de clientes por: nomeCliente (LIKE "%%term%%"), documento (= "term"), email (= "term")\n')
    cliente_vinc_reqs.add_run('RF-033:').bold = True
    cliente_vinc_reqs.add_run(' Sistema deve exibir lista paginada de clientes (20 por página) ordenada por nomeCliente ASC\n')
    cliente_vinc_reqs.add_run('RF-034:').bold = True
    cliente_vinc_reqs.add_run(' Auto-complete deve funcionar para campos cliente_fornecedor em lançamentos, retornando nomeCliente + documento\n')
    cliente_vinc_reqs.add_run('RF-035:').bold = True
    cliente_vinc_reqs.add_run(' Sistema deve validar integridade referencial: não permitir exclusão de cliente vinculado a OS/vendas/equipamentos ativos\n')
    cliente_vinc_reqs.add_run('RF-036:').bold = True
    cliente_vinc_reqs.add_run(' Relatório de clientes deve incluir contadores: total OS, total vendas, total equipamentos')

    # 4.4 Módulo de Produtos
    doc.add_heading('4.4 Módulo de Produtos', level=2)
    doc.add_paragraph('Gerencia cadastro e controle de produtos comercializados.')

    doc.add_heading('4.4.1 Cadastro de Produtos', level=3)
    produto_cad_reqs = doc.add_paragraph()
    produto_cad_reqs.add_run('RF-037:').bold = True
    produto_cad_reqs.add_run(' Sistema deve permitir cadastro de produtos com campos obrigatórios: descricao (varchar 80, única, não vazia), precoVenda (decimal 10,2 > 0), estoque (int >= 0)\n')
    produto_cad_reqs.add_run('RF-038:').bold = True
    produto_cad_reqs.add_run(' Campos opcionais: unidade (varchar 10), precoCompra (decimal 10,2), estoqueMinimo (int >= 0)\n')
    produto_cad_reqs.add_run('RF-039:').bold = True
    produto_cad_reqs.add_run(' Validação obrigatória: descricao.trim() não vazia, descricao única na tabela produtos, precoVenda > 0, estoque >= 0\n')
    produto_cad_reqs.add_run('RF-040:').bold = True
    produto_cad_reqs.add_run(' Validação condicional: se precoCompra informado então precoCompra > 0; se estoqueMinimo informado então estoqueMinimo >= 0\n')
    produto_cad_reqs.add_run('RF-041:').bold = True
    produto_cad_reqs.add_run(' Validação de formato: precoCompra e precoVenda devem aceitar apenas números decimais válidos (ex: 123.45 ou 123,45)\n')
    produto_cad_reqs.add_run('RF-042:').bold = True
    produto_cad_reqs.add_run(' Sistema deve permitir edição completa de todos os campos\n')
    produto_cad_reqs.add_run('RF-043:').bold = True
    produto_cad_reqs.add_run(' Sistema deve exibir lista paginada de produtos (20 por página) ordenada por descricao ASC\n')
    produto_cad_reqs.add_run('RF-044:').bold = True
    produto_cad_reqs.add_run(' Sistema deve permitir busca por descricao (LIKE "%%term%%")\n')
    produto_cad_reqs.add_run('RF-045:').bold = True
    produto_cad_reqs.add_run(' Sistema deve calcular margem de lucro: ((precoVenda - precoCompra) / precoCompra) * 100 quando precoCompra informado')

    doc.add_heading('4.4.2 Controle de Estoque', level=3)
    produto_estoque_reqs = doc.add_paragraph()
    produto_estoque_reqs.add_run('RF-046:').bold = True
    produto_estoque_reqs.add_run(' Sistema deve reduzir estoque automaticamente quando produto é adicionado a OS: UPDATE produtos SET estoque = estoque - quantidade WHERE idProdutos = produto_id\n')
    produto_estoque_reqs.add_run('RF-047:').bold = True
    produto_estoque_reqs.add_run(' Sistema deve reduzir estoque automaticamente quando produto é vendido: UPDATE produtos SET estoque = estoque - quantidade WHERE idProdutos = produto_id\n')
    produto_estoque_reqs.add_run('RF-048:').bold = True
    produto_estoque_reqs.add_run(' Sistema deve validar estoque suficiente antes de permitir adição em OS: SELECT estoque >= quantidade FROM produtos WHERE idProdutos = produto_id\n')
    produto_estoque_reqs.add_run('RF-049:').bold = True
    produto_estoque_reqs.add_run(' Sistema deve validar estoque suficiente antes de permitir venda: SELECT estoque >= quantidade FROM produtos WHERE idProdutos = produto_id\n')
    produto_estoque_reqs.add_run('RF-050:').bold = True
    produto_estoque_reqs.add_run(' Sistema deve exibir alerta visual (badge vermelho) quando estoque <= estoqueMinimo\n')
    produto_estoque_reqs.add_run('RF-051:').bold = True
    produto_estoque_reqs.add_run(' Sistema deve permitir entrada manual de estoque através de formulário específico: UPDATE produtos SET estoque = estoque + quantidade WHERE idProdutos = produto_id\n')
    produto_estoque_reqs.add_run('RF-052:').bold = True
    produto_estoque_reqs.add_run(' Sistema deve registrar histórico de movimentações de estoque (entrada/saída) em tabela separada\n')
    produto_estoque_reqs.add_run('RF-053:').bold = True
    produto_estoque_reqs.add_run(' Relatório de produtos deve incluir colunas calculadas: status_estoque (IF(estoque=0,"Esgotado",IF(estoque<=estoqueMinimo,"Baixo","Normal"))), valor_total_estoque (estoque * precoCompra)\n')
    produto_estoque_reqs.add_run('RF-054:').bold = True
    produto_estoque_reqs.add_run(' Sistema deve impedir edição de precoVenda se produto já foi vendido (verificar em itens_de_vendas)')

    # 4.5 Módulo de Serviços
    doc.add_heading('4.5 Módulo de Serviços', level=2)
    doc.add_paragraph('Gerencia cadastro de serviços prestados pela empresa.')

    doc.add_heading('4.5.1 Cadastro de Serviços', level=3)
    servico_cad_reqs = doc.add_paragraph()
    servico_cad_reqs.add_run('RF-055:').bold = True
    servico_cad_reqs.add_run(' Sistema deve permitir cadastro de serviços com campos obrigatórios: nome (varchar 45, único, não vazio), preco (decimal 10,2 > 0)\n')
    servico_cad_reqs.add_run('RF-056:').bold = True
    servico_cad_reqs.add_run(' Campo opcional: descricao (varchar 45)\n')
    servico_cad_reqs.add_run('RF-057:').bold = True
    servico_cad_reqs.add_run(' Validação obrigatória: nome.trim() não vazio, nome único na tabela servicos, preco > 0\n')
    servico_cad_reqs.add_run('RF-058:').bold = True
    servico_cad_reqs.add_run(' Validação de formato: preco deve aceitar apenas números decimais válidos (ex: 123.45 ou 123,45)\n')
    servico_cad_reqs.add_run('RF-059:').bold = True
    servico_cad_reqs.add_run(' Sistema deve permitir edição completa de todos os campos\n')
    servico_cad_reqs.add_run('RF-060:').bold = True
    servico_cad_reqs.add_run(' Sistema deve exibir lista paginada de serviços (20 por página) ordenada por nome ASC\n')
    servico_cad_reqs.add_run('RF-061:').bold = True
    servico_cad_reqs.add_run(' Sistema deve permitir busca por nome (LIKE "%%term%%") ou descricao (LIKE "%%term%%")\n')
    servico_cad_reqs.add_run('RF-062:').bold = True
    servico_cad_reqs.add_run(' Sistema deve validar integridade referencial: não permitir exclusão de serviço vinculado a OS ativas (verificar em servicos_os)')

    # 4.6 Módulo de Ordens de Serviço (OS)
    doc.add_heading('4.6 Módulo de Ordens de Serviço (OS)', level=2)
    doc.add_paragraph('Gerencia ordens de serviço para manutenção/reparo de equipamentos.')

    doc.add_heading('4.6.1 Cadastro de OS', level=3)
    os_cad_reqs = doc.add_paragraph()
    os_cad_reqs.add_run('RF-063:').bold = True
    os_cad_reqs.add_run(' Sistema deve permitir criar OS com campos obrigatórios: descricaoProduto (text, não vazio), defeito (text, não vazio), clientes_id (int, FK válido), usuarios_id (int, preenchido automaticamente), faturado (tinyint, inicia 0)\n')
    os_cad_reqs.add_run('RF-064:').bold = True
    os_cad_reqs.add_run(' Campos opcionais: dataInicial (date), dataFinal (date), garantia (varchar 45), status (varchar 45), observacoes (text), laudoTecnico (text), valorTotal (varchar 15), lancamento (int)\n')
    os_cad_reqs.add_run('RF-065:').bold = True
    os_cad_reqs.add_run(' Validação obrigatória: descricaoProduto.trim() não vazio, defeito.trim() não vazio, clientes_id existe na tabela clientes\n')
    os_cad_reqs.add_run('RF-066:').bold = True
    os_cad_reqs.add_run(' Campo dataInicial deve ser preenchido automaticamente com CURRENT_DATE na criação\n')
    os_cad_reqs.add_run('RF-067:').bold = True
    os_cad_reqs.add_run(' Campo usuarios_id deve ser preenchido automaticamente com session user ID\n')
    os_cad_reqs.add_run('RF-068:').bold = True
    os_cad_reqs.add_run(' Sistema deve permitir edição completa da OS enquanto faturado = 0\n')
    os_cad_reqs.add_run('RF-069:').bold = True
    os_cad_reqs.add_run(' Sistema deve exibir lista paginada de OS (20 por página) ordenada por dataInicial DESC\n')
    os_cad_reqs.add_run('RF-070:').bold = True
    os_cad_reqs.add_run(' Filtros devem incluir: status (exato), clientes_id (select), dataInicial BETWEEN periodo\n')
    os_cad_reqs.add_run('RF-071:').bold = True
    os_cad_reqs.add_run(' Status possíveis: "Orçamento", "Aprovado", "Em Andamento", "Aguardando Peças", "Finalizado", "Cancelado"\n')
    os_cad_reqs.add_run('RF-072:').bold = True
    os_cad_reqs.add_run(' Campo valorTotal deve ser calculado automaticamente como soma de todos subTotais de produtos_os + servicos_os')

    doc.add_heading('4.6.2 Equipamentos', level=3)
    os_equip_reqs = doc.add_paragraph()
    os_equip_reqs.add_run('RF-073:').bold = True
    os_equip_reqs.add_run(' Sistema deve permitir cadastro de equipamentos com campos obrigatórios: equipamento (varchar 150, não vazio), clientes_id (int, FK válido)\n')
    os_equip_reqs.add_run('RF-074:').bold = True
    os_equip_reqs.add_run(' Campos opcionais: num_serie (varchar 80), modelo (varchar 80), cor (varchar 45), descricao (varchar 150), tensao (varchar 45), potencia (varchar 45), voltagem (varchar 45), data_fabricacao (date), marcas_id (int)\n')
    os_equip_reqs.add_run('RF-075:').bold = True
    os_equip_reqs.add_run(' Validação obrigatória: equipamento.trim() não vazio, clientes_id existe na tabela clientes\n')
    os_equip_reqs.add_run('RF-076:').bold = True
    os_equip_reqs.add_run(' Sistema deve permitir vincular equipamentos existentes a OS através de tabela equipamentos_os\n')
    os_equip_reqs.add_run('RF-077:').bold = True
    os_equip_reqs.add_run(' Tabela equipamentos_os deve ter campos obrigatórios: equipamentos_id (int, FK), os_id (int, FK), campos opcionais: defeito_declarado (varchar 200), defeito_encontrado (varchar 200), solucao (varchar 45)\n')
    os_equip_reqs.add_run('RF-078:').bold = True
    os_equip_reqs.add_run(' Sistema deve permitir busca de equipamentos por cliente: SELECT * FROM equipamentos WHERE clientes_id = ? ORDER BY equipamento\n')
    os_equip_reqs.add_run('RF-079:').bold = True
    os_equip_reqs.add_run(' Sistema deve permitir cadastro de marcas através de tabela marcas com campos: marca (varchar 100, única), cadastro (date, auto), situacao (tinyint, default 1)\n')
    os_equip_reqs.add_run('RF-080:').bold = True
    os_equip_reqs.add_run(' Validação de num_serie único por equipamento quando informado')

    doc.add_heading('4.6.3 Produtos e Serviços em OS', level=3)
    os_prod_serv_reqs = doc.add_paragraph()
    os_prod_serv_reqs.add_run('RF-081:').bold = True
    os_prod_serv_reqs.add_run(' Sistema deve permitir adicionar produtos à OS através de tabela produtos_os com campos obrigatórios: quantidade (int > 0), os_id (int, FK), produtos_id (int, FK), subTotal (varchar 15, calculado)\n')
    os_prod_serv_reqs.add_run('RF-082:').bold = True
    os_prod_serv_reqs.add_run(' Sistema deve permitir adicionar serviços à OS através de tabela servicos_os com campos obrigatórios: os_id (int, FK), servicos_id (int, FK), subTotal (varchar 15, calculado)\n')
    os_prod_serv_reqs.add_run('RF-083:').bold = True
    os_prod_serv_reqs.add_run(' Validação obrigatória: produtos_id/servicos_id existem, quantidade > 0, estoque suficiente para produtos\n')
    os_prod_serv_reqs.add_run('RF-084:').bold = True
    os_prod_serv_reqs.add_run(' Cálculo subTotal produtos: quantidade * (SELECT precoVenda FROM produtos WHERE idProdutos = produtos_id)\n')
    os_prod_serv_reqs.add_run('RF-085:').bold = True
    os_prod_serv_reqs.add_run(' Cálculo subTotal serviços: 1 * (SELECT preco FROM servicos WHERE idServicos = servicos_id)\n')
    os_prod_serv_reqs.add_run('RF-086:').bold = True
    os_prod_serv_reqs.add_run(' Trigger UPDATE para recalcular valorTotal da OS: UPDATE os SET valorTotal = (SELECT SUM(subTotal) FROM produtos_os WHERE os_id = os.idOs) + (SELECT SUM(subTotal) FROM servicos_os WHERE os_id = os.idOs) WHERE idOs = os_id\n')
    os_prod_serv_reqs.add_run('RF-087:').bold = True
    os_prod_serv_reqs.add_run(' Sistema deve reduzir estoque automaticamente: UPDATE produtos SET estoque = estoque - quantidade WHERE idProdutos = produtos_id\n')
    os_prod_serv_reqs.add_run('RF-088:').bold = True
    os_prod_serv_reqs.add_run(' Sistema deve permitir remover itens da OS: DELETE FROM produtos_os WHERE id = ?; UPDATE produtos SET estoque = estoque + quantidade WHERE idProdutos = produtos_id\n')
    os_prod_serv_reqs.add_run('RF-089:').bold = True
    os_prod_serv_reqs.add_run(' Validação: não permitir remoção de itens se OS já faturada (faturado = 1)')

    doc.add_heading('4.6.4 Anexos e Finalização', level=3)
    os_anexos_reqs = doc.add_paragraph()
    os_anexos_reqs.add_run('RF-090:').bold = True
    os_anexos_reqs.add_run(' Sistema deve permitir anexar arquivos através de tabela anexos com campos obrigatórios: os_id (int, FK), campos opcionais: anexo (varchar 45), thumb (varchar 45), url (varchar 300), path (varchar 300)\n')
    os_anexos_reqs.add_run('RF-091:').bold = True
    os_anexos_reqs.add_run(' Sistema deve validar tipos MIME permitidos: image/*, application/pdf, application/msword, application/vnd.openxmlformats-officedocument.wordprocessingml.document\n')
    os_anexos_reqs.add_run('RF-092:').bold = True
    os_anexos_reqs.add_run(' Sistema deve limitar tamanho máximo de arquivo: 5MB por arquivo\n')
    os_anexos_reqs.add_run('RF-093:').bold = True
    os_anexos_reqs.add_run(' Faturamento deve executar: UPDATE os SET faturado = 1 WHERE idOs = ?; INSERT INTO lancamentos (valor, tipo, descricao, cliente_fornecedor, baixado, data_lancamento, id_usuario) VALUES (valorTotal, 1, "OS #" + idOs, cliente_nome, 1, CURRENT_DATE, user_id)\n')
    os_anexos_reqs.add_run('RF-094:').bold = True
    os_anexos_reqs.add_run(' Campo lancamento deve armazenar LAST_INSERT_ID() do lançamento financeiro criado\n')
    os_anexos_reqs.add_run('RF-095:').bold = True
    os_anexos_reqs.add_run(' OS faturada não deve permitir: edição de dados, adição/remoção de itens, alteração de status\n')
    os_anexos_reqs.add_run('RF-096:').bold = True
    os_anexos_reqs.add_run(' Sistema deve gerar PDF da OS com: dados da OS, dados do cliente, lista de produtos/serviços, valor total, anexos como links\n')
    os_anexos_reqs.add_run('RF-097:').bold = True
    os_anexos_reqs.add_run(' Validação de faturamento: OS deve ter pelo menos um item (produto ou serviço) para poder ser faturada')

    # 4.7 Módulo de Vendas
    doc.add_heading('4.7 Módulo de Vendas', level=2)
    doc.add_paragraph('Gerencia vendas de produtos para clientes.')

    doc.add_heading('4.7.1 Cadastro de Vendas', level=3)
    venda_cad_reqs = doc.add_paragraph()
    venda_cad_reqs.add_run('RF-098:').bold = True
    venda_cad_reqs.add_run(' Sistema deve permitir registrar vendas com campos obrigatórios: clientes_id (int, FK válido), usuarios_id (int, preenchido automaticamente), faturado (tinyint, inicia 0)\n')
    venda_cad_reqs.add_run('RF-099:').bold = True
    venda_cad_reqs.add_run(' Campos opcionais: dataVenda (date), valorTotal (varchar 45), desconto (varchar 45), lancamentos_id (int)\n')
    venda_cad_reqs.add_run('RF-100:').bold = True
    venda_cad_reqs.add_run(' Validação obrigatória: clientes_id existe na tabela clientes\n')
    venda_cad_reqs.add_run('RF-101:').bold = True
    venda_cad_reqs.add_run(' Campo dataVenda deve ser preenchido automaticamente com CURRENT_DATE\n')
    venda_cad_reqs.add_run('RF-102:').bold = True
    venda_cad_reqs.add_run(' Campo usuarios_id deve ser preenchido automaticamente com session user ID\n')
    venda_cad_reqs.add_run('RF-103:').bold = True
    venda_cad_reqs.add_run(' Sistema deve permitir edição completa da venda enquanto faturado = 0\n')
    venda_cad_reqs.add_run('RF-104:').bold = True
    venda_cad_reqs.add_run(' Sistema deve exibir lista paginada de vendas (20 por página) ordenada por dataVenda DESC\n')
    venda_cad_reqs.add_run('RF-105:').bold = True
    venda_cad_reqs.add_run(' Filtros devem incluir: clientes_id (select), dataVenda BETWEEN periodo, faturado (0/1)\n')
    venda_cad_reqs.add_run('RF-106:').bold = True
    venda_cad_reqs.add_run(' Validação de desconto: deve ser <= valorTotal quando informado\n')
    venda_cad_reqs.add_run('RF-107:').bold = True
    venda_cad_reqs.add_run(' Campo valorTotal deve ser calculado como: soma de todos subTotais de itens_de_vendas - desconto')

    doc.add_heading('4.7.2 Itens de Venda', level=3)
    venda_itens_reqs = doc.add_paragraph()
    venda_itens_reqs.add_run('RF-108:').bold = True
    venda_itens_reqs.add_run(' Sistema deve permitir adicionar itens através de tabela itens_de_vendas com campos obrigatórios: quantidade (int > 0), vendas_id (int, FK), produtos_id (int, FK), subTotal (varchar 45, calculado)\n')
    venda_itens_reqs.add_run('RF-109:').bold = True
    venda_itens_reqs.add_run(' Validação obrigatória: produtos_id existe na tabela produtos, quantidade > 0, estoque suficiente (SELECT estoque >= quantidade FROM produtos WHERE idProdutos = produtos_id)\n')
    venda_itens_reqs.add_run('RF-110:').bold = True
    venda_itens_reqs.add_run(' Cálculo subTotal: quantidade * (SELECT precoVenda FROM produtos WHERE idProdutos = produtos_id)\n')
    venda_itens_reqs.add_run('RF-111:').bold = True
    venda_itens_reqs.add_run(' Trigger UPDATE para recalcular valorTotal: UPDATE vendas SET valorTotal = (SELECT SUM(subTotal) FROM itens_de_vendas WHERE vendas_id = vendas.idVendas) - COALESCE(desconto, 0) WHERE idVendas = vendas_id\n')
    venda_itens_reqs.add_run('RF-112:').bold = True
    venda_itens_reqs.add_run(' Sistema deve reduzir estoque automaticamente: UPDATE produtos SET estoque = estoque - quantidade WHERE idProdutos = produtos_id\n')
    venda_itens_reqs.add_run('RF-113:').bold = True
    venda_itens_reqs.add_run(' Sistema deve permitir remover itens: DELETE FROM itens_de_vendas WHERE idItens = ?; UPDATE produtos SET estoque = estoque + quantidade WHERE idProdutos = produtos_id\n')
    venda_itens_reqs.add_run('RF-114:').bold = True
    venda_itens_reqs.add_run(' Validação: não permitir remoção se venda já faturada (faturado = 1)\n')
    venda_itens_reqs.add_run('RF-115:').bold = True
    venda_itens_reqs.add_run(' Validação: venda deve ter pelo menos um item para ser salva\n')
    venda_itens_reqs.add_run('RF-116:').bold = True
    venda_itens_reqs.add_run(' Sistema deve impedir venda de produto com estoque = 0')

    doc.add_heading('4.7.3 Faturamento e Finalização', level=3)
    venda_fat_reqs = doc.add_paragraph()
    venda_fat_reqs.add_run('RF-117:').bold = True
    venda_fat_reqs.add_run(' Faturamento deve executar: UPDATE vendas SET faturado = 1 WHERE idVendas = ?; INSERT INTO lancamentos (valor, tipo, descricao, cliente_fornecedor, baixado, data_lancamento, id_usuario) VALUES (valorTotal, 1, "Venda #" + idVendas, cliente_nome, 1, CURRENT_DATE, user_id)\n')
    venda_fat_reqs.add_run('RF-118:').bold = True
    venda_fat_reqs.add_run(' Campo lancamentos_id deve armazenar LAST_INSERT_ID() do lançamento financeiro\n')
    venda_fat_reqs.add_run('RF-119:').bold = True
    venda_fat_reqs.add_run(' Venda faturada não deve permitir: edição de dados, adição/remoção de itens, alteração de desconto\n')
    venda_fat_reqs.add_run('RF-120:').bold = True
    venda_fat_reqs.add_run(' Sistema deve permitir cancelamento de venda não faturada: UPDATE vendas SET status = 0 WHERE idVendas = ? AND faturado = 0; UPDATE produtos SET estoque = estoque + quantidade FROM itens_de_vendas WHERE vendas_id = ?\n')
    venda_fat_reqs.add_run('RF-121:').bold = True
    venda_fat_reqs.add_run(' Validação de faturamento: venda deve ter pelo menos um item e valorTotal > 0\n')
    venda_fat_reqs.add_run('RF-122:').bold = True
    venda_fat_reqs.add_run(' Sistema deve gerar comprovante de venda em PDF com: dados da venda, dados do cliente, lista de itens, valor total, data/hora')

    # Requisitos Não Funcionais
    doc.add_heading('5. Requisitos Não Funcionais', level=1)

    # Segurança
    doc.add_heading('5.1 Segurança', level=2)
    security_reqs = doc.add_paragraph()
    security_reqs.add_run('RNF-001:').bold = True
    security_reqs.add_run(' Autenticação obrigatória para acesso ao sistema\n')
    security_reqs.add_run('RNF-002:').bold = True
    security_reqs.add_run(' Controle de permissões por usuário\n')
    security_reqs.add_run('RNF-003:').bold = True
    security_reqs.add_run(' Criptografia de dados sensíveis (senhas, números de cartão)\n')
    security_reqs.add_run('RNF-004:').bold = True
    security_reqs.add_run(' Proteção contra SQL Injection\n')
    security_reqs.add_run('RNF-005:').bold = True
    security_reqs.add_run(' Proteção contra XSS\n')
    security_reqs.add_run('RNF-006:').bold = True
    security_reqs.add_run(' Sessões armazenadas em banco de dados\n')
    security_reqs.add_run('RNF-007:').bold = True
    security_reqs.add_run(' Controle de acesso baseado em roles')

    # Performance
    doc.add_heading('5.2 Performance', level=2)
    perf_reqs = doc.add_paragraph()
    perf_reqs.add_run('RNF-008:').bold = True
    perf_reqs.add_run(' Tempo de resposta < 2 segundos para operações comuns\n')
    perf_reqs.add_run('RNF-009:').bold = True
    perf_reqs.add_run(' Suporte a múltiplos usuários simultâneos\n')
    perf_reqs.add_run('RNF-010:').bold = True
    perf_reqs.add_run(' Paginação em listagens grandes\n')
    perf_reqs.add_run('RNF-011:').bold = True
    perf_reqs.add_run(' Cache de consultas frequentes\n')
    perf_reqs.add_run('RNF-012:').bold = True
    perf_reqs.add_run(' Otimização de queries SQL')

    # Usabilidade
    doc.add_heading('5.3 Usabilidade', level=2)
    usability_reqs = doc.add_paragraph()
    usability_reqs.add_run('RNF-013:').bold = True
    usability_reqs.add_run(' Interface responsiva (desktop/mobile)\n')
    usability_reqs.add_run('RNF-014:').bold = True
    usability_reqs.add_run(' Navegação intuitiva\n')
    usability_reqs.add_run('RNF-015:').bold = True
    usability_reqs.add_run(' Auto-complete em campos de busca\n')
    usability_reqs.add_run('RNF-016:').bold = True
    usability_reqs.add_run(' Validação em tempo real de formulários\n')
    usability_reqs.add_run('RNF-017:').bold = True
    usability_reqs.add_run(' Feedback visual para ações do usuário\n')
    usability_reqs.add_run('RNF-018:').bold = True
    usability_reqs.add_run(' Design consistente')

    # Requisitos de Dados
    doc.add_heading('6. Requisitos de Dados', level=1)

    # Modelo de Dados
    doc.add_heading('6.1 Modelo de Dados', level=2)
    doc.add_paragraph(
        'O sistema utiliza MySQL/MariaDB com as seguintes tabelas principais:'
    )

    tables = [
        'usuarios - Dados dos usuários do sistema',
        'permissoes - Controle de permissões',
        'clientes - Cadastro de clientes',
        'produtos - Cadastro de produtos',
        'servicos - Cadastro de serviços',
        'lancamentos - Registros financeiros',
        'cartoes - Cartões de crédito',
        'faturas - Faturas mensais',
        'lancamentos_faturas - Lançamentos das faturas',
        'lancamentos_faturas_assoc - Associação lançamentos-faturas',
        'os - Ordens de serviço',
        'produtos_os - Produtos em OS',
        'servicos_os - Serviços em OS',
        'vendas - Registros de vendas',
        'itens_de_vendas - Itens das vendas',
        'anexos - Arquivos anexados',
        'logs - Log de atividades'
    ]

    for table in tables:
        doc.add_paragraph(f'• {table}', style='List Bullet')

    # Regras de Negócio
    doc.add_heading('7. Regras de Negócio', level=1)

    # Regras Financeiras
    doc.add_heading('7.1 Regras Financeiras', level=2)
    financial_rules = doc.add_paragraph()
    financial_rules.add_run('Regras de Negócio Financeiras:').bold = True
    financial_rules.add_run("""
• Valores negativos representam despesas/saídas
• Valores positivos representam receitas/entradas
• Lançamentos parcelados geram múltiplas associações com faturas
• Faturas fechadas não permitem novos lançamentos
• Compras de terceiros geram pendências adicionais
• Vencimento de faturas baseado no dia configurado por cartão
• Saldo disponível = saldo provisório - pendências
• Lançamentos ocultos não aparecem na visualização padrão
""")

    # Regras de Estoque
    doc.add_heading('7.2 Regras de Estoque', level=2)
    stock_rules = doc.add_paragraph()
    stock_rules.add_run('Regras de Estoque:').bold = True
    stock_rules.add_run("""
• Estoque atualizado automaticamente ao adicionar produtos em OS
• Alerta quando estoque abaixo do mínimo
• Controle de entrada/saída de produtos
• Preços de compra e venda independentes
""")

    # Regras de OS
    doc.add_heading('7.3 Regras de Ordens de Serviço', level=2)
    os_rules = doc.add_paragraph()
    os_rules.add_run('Regras de OS:').bold = True
    os_rules.add_run("""
• OS faturada não pode ser alterada
• Produtos adicionados reduzem estoque automaticamente
• Valor total calculado automaticamente
• Status controla fluxo de trabalho
• Anexos limitados a tipos específicos
""")

    # Fluxos de Processo
    doc.add_heading('8. Fluxos de Processo', level=1)

    # Fluxo de Lançamentos
    doc.add_heading('8.1 Fluxo de Lançamentos Financeiros', level=2)
    lanc_flow = doc.add_paragraph()
    lanc_flow.add_run('Fluxo de Lançamentos:').bold = True
    lanc_flow.add_run("""
1. Usuário acessa módulo de lançamentos
2. Seleciona tipo (entrada/saída)
3. Preenche dados (descrição, valor, data, fornecedor)
4. Opcionalmente marca como pago e seleciona forma de pagamento
5. Sistema valida dados e salva
6. Atualiza saldos automaticamente
""")

    # Fluxo de Faturas
    doc.add_heading('8.2 Fluxo de Faturas', level=2)
    fatura_flow = doc.add_paragraph()
    fatura_flow.add_run('Fluxo de Faturas:').bold = True
    fatura_flow.add_run("""
1. Sistema cria faturas automaticamente baseado em configuração
2. Usuário registra compras no cartão
3. Sistema associa lançamentos à fatura correspondente
4. No vencimento, usuário paga fatura
5. Sistema marca fatura como paga
6. Opcionalmente vincula ao módulo financeiro
""")

    # Fluxo de OS
    doc.add_heading('8.3 Fluxo de Ordens de Serviço', level=2)
    os_flow = doc.add_paragraph()
    os_flow.add_run('Fluxo de OS:').bold = True
    os_flow.add_run("""
1. Cliente solicita serviço
2. Técnico cria OS com dados do equipamento
3. Adiciona produtos e serviços necessários
4. Sistema calcula valor total
5. OS é executada e finalizada
6. Cliente aprova e paga
7. Sistema gera lançamento financeiro
""")

    # Interfaces do Usuário
    doc.add_heading('9. Interfaces do Usuário', level=1)

    # Dashboard
    doc.add_heading('9.1 Dashboard Financeiro', level=2)
    dashboard_desc = doc.add_paragraph()
    dashboard_desc.add_run('Dashboard Principal:').bold = True
    dashboard_desc.add_run("""
• Saldo disponível em conta
• Saldo de entradas/saídas pendentes
• Posição consolidada do período
• Lançamentos recentes
• Faturas em aberto
• Gráficos de evolução financeira
""")

    # Formulários
    doc.add_heading('9.2 Formulários', level=2)
    forms_desc = doc.add_paragraph()
    forms_desc.add_run('Principais Formulários:').bold = True
    forms_desc.add_run("""
• Cadastro/edição de lançamentos financeiros
• Cadastro de cartões e configuração de faturas
• Cadastro de clientes, produtos, serviços
• Criação e edição de ordens de serviço
• Filtros e pesquisas avançadas
• Relatórios financeiros
""")

    # Relatórios
    doc.add_heading('10. Relatórios', level=1)
    reports_desc = doc.add_paragraph()
    reports_desc.add_run('Relatórios Disponíveis:').bold = True
    reports_desc.add_run("""
• Extrato de lançamentos por período
• Relatório de faturas
• Relatório de clientes
• Relatório de produtos
• Relatório de serviços
• Relatório de OS
• Relatório financeiro consolidado
• Gráficos de evolução
""")

    # Conclusão
    doc.add_heading('11. Conclusão', level=1)
    doc.add_paragraph(
        'Este documento apresenta uma visão completa dos requisitos do Sistema de Gestão Financeira Pessoal, '
        'baseada na análise detalhada do código fonte. O sistema oferece funcionalidades abrangentes para '
        'controle financeiro pessoal, gestão de clientes, produtos, serviços e ordens de serviço, '
        'atendendo às necessidades de usuários que desejam ter controle total sobre suas finanças e operações.'
    )

    # Salvar documento
    doc.save('docs/documento_requisitos.docx')
    print("Documento de requisitos criado com sucesso: docs/documento_requisitos.docx")

if __name__ == "__main__":
    create_requirements_document()